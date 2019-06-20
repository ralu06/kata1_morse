from docx import Document
from docx.shared import Pt
from time import strftime, gmtime

document = Document()

document.add_heading('Telegrama', 0)

hoy = strftime("%d - %b - %Y", gmtime())

parrafoFecha = document.add_paragraph(hoy) 
parrafoFecha.alignment = 2

de = document.add_paragraph()
de.add_run('From:').bold = True 
de.add_run(' Remitente')

para = document.add_paragraph()
para.add_run('To:').bold = True 
para.add_run(' Destinatario')

document.add_heading('Mensaje', level=1)

table = document.add_table(rows=2, cols=3)
table.style =  'LightShading'

table.style.font.name = 'Courier'
table.style.font.size = Pt(12)
# Variable Local
# style = table.style
# font = style.font
# font.name = 'Courrier'
# font.size = Pt(12)

celda_Morse = table.rows[0].cells[0]
celda_Morse.text = '.-.-.'

celda_plano = table.rows[1].cells[0]
celda_plano.text = 'Hola'

document.save('demo.docx')